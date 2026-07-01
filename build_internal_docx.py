#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate internal standard docx from the single combined markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent
REF = BASE / "_archive" / "参考来源"
TEMPLATE = REF / "参考-云平台-机器人通信协议V1.0.0-标准版（IOT基线）.docx"
MD_FILE = REF / "归档-云平台-机器人通信协议-内部标准版-R13.md"
OUT_FILE = BASE / "云平台-机器人通信协议-内部标准版.docx"
OUT_ALT = BASE / "云平台-机器人通信协议-内部标准版-生成.docx"

YELLOW = "FFF2CC"
HEADER_FILL = "D9E2F3"
FONT = "微软雅黑"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style_name = style_map.get(level, "Heading 3")
    sizes = {1: 14, 2: 12, 3: 11, 4: 10.5}
    spacings = {
        1: (Pt(12), Pt(6)),
        2: (Pt(8), Pt(4)),
        3: (Pt(6), Pt(3)),
        4: (Pt(4), Pt(2)),
    }
    available = {s.name for s in doc.styles}
    p = doc.add_paragraph(style=style_name) if style_name in available else doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, sizes.get(level, 11), True)
    before, after = spacings.get(level, (Pt(6), Pt(3)))
    p.paragraph_format.space_before = before
    p.paragraph_format.space_after = after
    return p


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def add_code(doc, text):
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        p.paragraph_format.left_indent = Cm(0.6)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        shd.set(qn("w:val"), "clear")
        p._element.get_or_add_pPr().append(shd)


def is_ext_row(row: list[str]) -> bool:
    joined = "".join(row)
    return any(x in joined for x in ("P1", "新增", "修改"))


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        ext = is_ext_row(row) and i > 0
        for j in range(cols):
            txt = row[j].strip() if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            set_run_font(run, 10, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, HEADER_FILL)
            elif ext:
                set_cell_shading(cell, YELLOW)
    doc.add_paragraph()


def parse_md(path: Path) -> list[dict]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code)})
            i += 1
            continue
        if "|" in line and line.strip().startswith("|"):
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                if re.match(r"^\|\s*[-:]+", lines[i].strip()):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
        elif line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
        elif line.startswith("#### "):
            blocks.append({"type": "h4", "text": line[5:].strip()})
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            blocks.append({"type": "bold", "text": line.strip().strip("*")})
        elif line.strip():
            blocks.append({"type": "para", "text": line.strip()})
        i += 1
    return blocks


def init_doc_from_template():
    if TEMPLATE.exists():
        doc = Document(str(TEMPLATE))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag in ("p", "tbl"):
                body.remove(child)
        return doc
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def build():
    blocks = parse_md(MD_FILE)
    doc = init_doc_from_template()

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run("云平台-机器人通信协议")
    set_run_font(r1, 16, True)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("V1.0.0（内部标准版）")
    set_run_font(r2, 14, True)
    doc.add_paragraph()

    skip_title = True
    for b in blocks:
        if b["type"] == "h1":
            if skip_title:
                skip_title = False
                continue
            add_heading(doc, b["text"], 1)
        elif b["type"] == "h2":
            add_heading(doc, b["text"], 1)
        elif b["type"] == "h3":
            add_heading(doc, b["text"], 2)
        elif b["type"] == "h4":
            add_heading(doc, b["text"], 4)
        elif b["type"] == "bold":
            add_para(doc, b["text"], bold=True)
        elif b["type"] == "para":
            add_para(doc, b["text"])
        elif b["type"] == "code":
            add_code(doc, b["text"])
        elif b["type"] == "table":
            add_table(doc, b["rows"])

    try:
        doc.save(str(OUT_FILE))
        print(f"OK: {OUT_FILE}")
    except PermissionError:
        doc.save(str(OUT_ALT))
        print(f"OK (原文件被占用，已另存): {OUT_ALT}")


if __name__ == "__main__":
    build()
