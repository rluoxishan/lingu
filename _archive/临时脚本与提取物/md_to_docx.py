#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert project markdown protocol docs to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_default_font(doc: Document, name: str = "微软雅黑", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        shading = p._element.get_or_add_pPr()
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F5F5F5")
        shading.append(shd)


def parse_inline(text: str):
    """Yield (text, bold) segments for simple **bold** markup."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True
        elif part:
            yield part, False


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    for segment, bold in parse_inline(text):
        run = p.add_run(segment)
        run.bold = bold


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            for segment, bold in parse_inline(cell_text.strip()):
                run = p.add_run(segment)
                run.bold = bold
                run.font.size = Pt(10)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    set_default_font(doc)

    # Title from first H1
    title = md_path.stem
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(18)

    doc.add_paragraph()

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table = False
    skip_first_h1 = True

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            # skip separator row
            if re.match(r"^\|\s*[-:]+", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            # peek: end table if next line not table
            if i >= len(lines) or "|" not in lines[i] or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                doc.add_paragraph()
                in_table = False
                table_rows = []
            continue

        in_table = False

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            if skip_first_h1:
                skip_first_h1 = False
            else:
                doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:])
        elif stripped.startswith("- [ ] "):
            add_rich_paragraph(doc, "☐ " + stripped[6:], style="List Bullet")
        elif stripped.startswith("- "):
            add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            add_rich_paragraph(doc, stripped)

        i += 1

    doc.save(str(docx_path))
    print(f"OK: {docx_path}")


def main(argv: list[str]) -> int:
    base = Path(__file__).resolve().parent
    files = argv[1:] if len(argv) > 1 else [
        "云平台-机器人通信协议-内部标准版.md",
        "灵鱿客户对接差异说明.md",
    ]
    for name in files:
        md = base / name
        if not md.exists():
            print(f"SKIP (not found): {md}")
            continue
        out = md.with_suffix(".docx")
        md_to_docx(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
