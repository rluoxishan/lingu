# -*- coding: utf-8 -*-
"""生成《发给Hasun-监控页雷达感知API确认单.docx》"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "发给Hasun-监控页雷达感知API确认单.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.get_or_add_shd()
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(12)


def add_checkbox_items(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    title = doc.add_heading("云平台远程监控 — 室内雷达感知 API 确认单", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "请 Hasun/云平台填写。用于监控页室内车型雷达感知图（PGM 底图 + 1010003 障碍物叠加）联调定稿。"
    )
    meta = doc.add_paragraph()
    meta.add_run("协议依据：").bold = True
    meta.add_run("《云平台-机器人通信协议-内部标准版》§6.4 1010003、§7.12 2010012、§7.10 2010010")
    meta2 = doc.add_paragraph()
    meta2.add_run("前端状态：").bold = True
    meta2.add_run("useMonitorHighFreq.ts / IndoorMapPanel 已起草，待云平台 API 就绪后联调。")
    meta3 = doc.add_paragraph()
    meta3.add_run("参考样例地图：").bold = True
    meta3.add_run("remote-monitor/地图与雷达文件/lhgk_101.{pgm,yaml,json}")
    meta4 = doc.add_paragraph()
    meta4.add_run("文档版本：").bold = True
    meta4.add_run("2026-07-04")

    doc.add_heading("转发话术（可复制发微信）", level=1)
    add_code(
        doc,
        "Hasun 你好，室内监控雷达感知我们在前端起草了 API 契约，需要云平台补 3 个能力：\n\n"
        "1. POST /device/instructions 支持 type=2010012（开/关 1010003 高频）\n"
        "2. GET /device/map_meta — 返回 PGM 转 PNG + origin/resolution（参考 lhgk_101.yaml）\n"
        "3. WebSocket /device/perception/stream?deviceId=xxx 推送 1010003 解析体\n"
        "   （或 GET /device/perception/latest 轮询兜底）\n\n"
        "室内 MVP 每帧最少：mapId、position_xyz、heading、obstacles[].polygon。\n"
        "详细字段表见附件 Word 文档各章节。",
    )

    doc.add_heading("1. 环境与联调信息（请填写）", level=1)
    add_table(
        doc,
        ["项", "填写"],
        [
            ["测试环境 Base URL", "https://sztu.lingubot.cn/admin-api"],
            ["WebSocket 地址", "wss://…/admin-api/device/perception/stream?deviceId=&token="],
            ["测试 indoor deviceId", ""],
            ["deviceId 与 MQTT clientId 映射", ""],
            ["设备绑定 mapId", "例：lhgk_101"],
            ["position_xyz 与 PGM 是否同坐标系", "☐ 已确认  ☐ 待确认"],
        ],
    )

    doc.add_heading("2. 数据链路说明", level=1)
    doc.add_paragraph(
        "监控页打开（室内车，无 GPS）→ POST /device/instructions type=2010012 status=ON "
        "→ 车端 dev/pub 1010003 @10Hz → 云平台订阅 EMQX 缓存 → WS 推前端 → PGM 底图 + obstacles 叠加。"
    )
    p = doc.add_paragraph()
    run = p.add_run("重要：1010001 / select_device_detail 不能替代 1010003（无 obstacles、仅 1Hz）。")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_heading("3. API 一览", level=1)
    add_table(
        doc,
        ["#", "方法", "路径", "状态", "说明"],
        [
            ["A", "POST", "/device/instructions", "已有，需支持 2010012", "开/关 1010003"],
            ["B", "GET", "/device/map_meta", "待实现", "室内 PGM 元数据 + 图片 URL"],
            ["C", "WS", "/device/perception/stream", "待实现", "推送 1010003 解析体"],
            ["D", "GET", "/device/perception/latest", "待实现", "WS 兜底，最新一帧"],
        ],
    )

    doc.add_heading("4. A — 2010012 高频开关", level=1)
    doc.add_paragraph("路径：POST /device/instructions")
    doc.add_heading("4.1 开启（监控页打开时）", level=2)
    add_code(
        doc,
        '{\n'
        '  "deviceId": "lingu_indoor_01",\n'
        '  "type": "2010012",\n'
        '  "data": {\n'
        '    "status": "ON",\n'
        '    "sensorTypes": ["RTK", "OBSTACLE", "TRAJECTORY"],\n'
        '    "frequencyHz": 10,\n'
        '    "durationSec": 300,\n'
        '    "keepAlivePeriods": 3,\n'
        '    "heartbeatIntervalSec": 5\n'
        '  }\n'
        "}",
    )
    doc.add_heading("4.2 关闭（监控页关闭时）", level=2)
    add_code(
        doc,
        '{\n'
        '  "deviceId": "lingu_indoor_01",\n'
        '  "type": "2010012",\n'
        '  "data": { "status": "OFF" }\n'
        "}",
    )
    doc.add_paragraph("前端每 5 秒重发 ON 作为心跳（与 heartbeatIntervalSec 一致）。")
    doc.add_heading("4.3 sensorTypes 与雷达图", level=2)
    add_table(
        doc,
        ["sensorTypes", "1010003 字段", "雷达图用途"],
        [
            ["RTK", "position_xyz, heading, speedMps", "车位姿（必须）"],
            ["OBSTACLE", "obstacles[]", "障碍物多边形（必须）"],
            ["TRAJECTORY", "trajectoryPoints[]", "局部规划线（建议）"],
            ["ULTRASONIC", "ultrasonicSense[]", "超声（可选）"],
        ],
    )
    doc.add_paragraph("室内 MVP 建议开启：RTK + OBSTACLE + TRAJECTORY。")

    doc.add_heading("5. B — GET /device/map_meta", level=1)
    doc.add_paragraph("请求：GET /device/map_meta?deviceId={deviceId}&mapId={mapId?}")
    doc.add_paragraph("不传 mapId 时按设备当前绑定地图返回；需与 1010003 帧内 mapId 一致。")
    doc.add_heading("5.1 响应 data 字段", level=2)
    add_table(
        doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["mapId", "String", "是", "如 lhgk_101"],
            ["imageUrl", "String", "是", "PNG/JPG URL（PGM 转换）"],
            ["width", "Integer", "是", "像素宽"],
            ["height", "Integer", "是", "像素高"],
            ["resolution", "Number", "是", "米/像素，如 0.05"],
            ["origin", "Number[3]", "是", "图像左下角世界坐标 [x,y,yaw]"],
            ["negate", "Integer", "否", "0=白空闲黑占用"],
        ],
    )
    doc.add_heading("5.2 响应示例", level=2)
    add_code(
        doc,
        '{\n'
        '  "code": 0,\n'
        '  "data": {\n'
        '    "mapId": "lhgk_101",\n'
        '    "imageUrl": "https://sztu.lingubot.cn/static/maps/lhgk_101.png",\n'
        '    "width": 3153,\n'
        '    "height": 944,\n'
        '    "resolution": 0.05,\n'
        '    "origin": [-11.0641, -37.6333, 0.0]\n'
        "  }\n"
        "}",
    )
    doc.add_heading("5.3 坐标换算", level=2)
    add_code(
        doc,
        "pixelX = (worldX - origin[0]) / resolution\n"
        "pixelY = height - (worldY - origin[1]) / resolution",
    )

    doc.add_heading("6. C — WebSocket /device/perception/stream", level=1)
    add_code(
        doc,
        "wss://{host}/admin-api/device/perception/stream?deviceId={deviceId}&token={accessToken}",
    )
    doc.add_paragraph("鉴权：Query token 或 Header Authorization（请确认一种）。仅推送该 deviceId 的 1010003。")
    doc.add_heading("6.1 下行消息示例", level=2)
    add_code(
        doc,
        '{\n'
        '  "code": 0,\n'
        '  "msg": "1010003",\n'
        '  "data": {\n'
        '    "mapId": "lhgk_101",\n'
        '    "position_xyz": "64.882,-14.310,0",\n'
        '    "heading": 92.5,\n'
        '    "speedMps": 0.35,\n'
        '    "trajectoryPoints": [{ "x": 65.0, "y": -14.5, "z": 0 }],\n'
        '    "obstacles": [{\n'
        '      "id": 1, "type": "PEDESTRIAN",\n'
        '      "polygon": [\n'
        '        { "x": 68.1, "y": -12.0 }, { "x": 68.5, "y": -12.0 },\n'
        '        { "x": 68.5, "y": -11.6 }, { "x": 68.1, "y": -11.6 }\n'
        "      ]\n"
        "    }]\n"
        "  }\n"
        "}",
    )
    doc.add_heading("6.2 data 字段（室内 MVP）", level=2)
    add_table(
        doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["mapId", "String", "是", "与 map_meta 一致"],
            ["position_xyz", "String", "是", "x,y,z  map 坐标系（米）"],
            ["heading", "Number", "是", "航向角（度）"],
            ["obstacles", "Array", "有感知时必填", "障碍物多边形，map 坐标"],
            ["trajectoryPoints", "Array", "否", "局部规划轨迹"],
            ["ultrasonicSense", "Array", "否", "distance 单位厘米"],
        ],
    )
    doc.add_heading("6.3 obstacles[] 项", level=2)
    add_table(
        doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["id", "Integer", "是", "跟踪 id"],
            ["type", "String", "是", "PEDESTRIAN / VEHICLE / STATIC / UNKNOWN"],
            ["polygon", "Array", "是", "≥3 顶点 {x,y}，map 坐标（米）"],
        ],
    )

    doc.add_heading("7. D — GET /device/perception/latest", level=1)
    doc.add_paragraph("请求：GET /device/perception/latest?deviceId={deviceId}")
    doc.add_paragraph("响应 data 与 WebSocket 相同；无数据时 data=null。前端 WS 失败 4 次后改 200ms 轮询。")

    doc.add_heading("8. 车端 MQTT 1010003（云平台订阅来源）", level=1)
    doc.add_paragraph("Topic：dev/pub/{clientId}，QoS 0，type=1010003。")
    add_code(
        doc,
        '{\n'
        '  "type": "1010003",\n'
        '  "data": {\n'
        '    "mapId": "lhgk_101",\n'
        '    "position_xyz": "64.882,-14.310,0",\n'
        '    "heading": 92.5,\n'
        '    "obstacles": []\n'
        "  }\n"
        "}",
    )
    doc.add_paragraph("云平台：订阅 → 解析 → Redis 最新帧 → 推 WS。")

    doc.add_heading("9. 验收清单 — 云平台（Hasun）", level=1)
    add_checkbox_items(
        doc,
        [
            "POST /device/instructions 支持 2010012，转发至 dev/sub/{clientId}",
            "订阅 dev/pub/+，识别 type=1010003",
            "实现 GET /device/map_meta（至少 lhgk_101）",
            "实现 WS /device/perception/stream 或 GET /device/perception/latest",
            "确认 position_xyz 与 PGM origin/resolution 同一 map 系",
            "监控页关闭后会话结束，车端停止 1010003",
        ],
    )

    doc.add_heading("10. 验收清单 — 车端", level=1)
    add_checkbox_items(
        doc,
        [
            "收到 2010012 ON 后 1010003 @10Hz",
            "每帧含 mapId、position_xyz、heading",
            "有障碍时 obstacles[].polygon 为 map 世界坐标（米）",
        ],
    )

    doc.add_heading("11. 备注与答复区", level=1)
    add_table(
        doc,
        ["问题", "Hasun 答复"],
        [
            ["2010012 是否已支持？", ""],
            ["map_meta 预计上线时间？", ""],
            ["perception/stream 鉴权方式？", ""],
            ["联调 indoor deviceId？", ""],
            ["其他说明", ""],
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
