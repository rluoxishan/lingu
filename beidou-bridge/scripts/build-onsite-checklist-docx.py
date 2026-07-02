#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate onsite confirmation checklist Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "到场确认清单.docx"
FONT = "微软雅黑"
HEADER_FILL = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 16, True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10, False)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 2")
    available = {s.name for s in doc.styles}
    p = doc.add_paragraph(style=style) if style in available else doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, {1: 14, 2: 12, 3: 11}.get(level, 11), True)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], HEADER_FILL)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, 10, True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, 10, False)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_title(doc, "beidou-bridge 到场确认清单")
    add_subtitle(
        doc,
        "项目：云平台 ↔ 中转 ↔ 北斗  |  现场车：LU2606000100  |  验收：B1 读推送  |  版本：2026-07-02",
    )

    add_para(
        doc,
        "用途：中转机、北斗环境不确定时，到场边联调边填写。程序操作见《现场安装与真北斗联调教程》。",
    )

    add_heading(doc, "〇、出发前 U 盘必备（有网电脑上下载好）", 1)
    add_para(
        doc,
        "中转机若无外网，下列文件必须在出发前拷入 U 盘/移动硬盘；zip 已含 node_modules，到场无需 npm install。",
        bold=True,
    )
    add_table(
        doc,
        ["#", "文件名", "在本机路径 / 下载方式", "大小约", "□已拷"],
        [
            [
                "1",
                "beidou-bridge-site-20260701.zip",
                "beidou-bridge\\release\\ 目录",
                "约 2～5 MB",
                "□",
            ],
            [
                "2",
                ".env",
                "beidou-bridge 根目录（单独拷，勿公开）",
                "很小",
                "□",
            ],
            [
                "3",
                "node-v20.x-x64.msi",
                "https://nodejs.org/ 下载 LTS Windows 64-bit 安装包",
                "约 30 MB",
                "□",
            ],
            [
                "4",
                "到场确认清单.docx",
                "beidou-bridge\\docs\\",
                "很小",
                "□",
            ],
            [
                "5",
                "现场安装与真北斗联调教程.md",
                "beidou-bridge\\docs\\（或打印）",
                "很小",
                "□",
            ],
            [
                "6",
                "给北斗方-接口一页纸.txt",
                "beidou-bridge\\docs\\",
                "很小",
                "□",
            ],
        ],
        [0.8, 4.2, 6.5, 2, 1.2],
    )
    add_para(doc, "可选（非必须）：")
    add_table(
        doc,
        ["项", "说明", "□"],
        [
            ["CHECK-env.bat 等 bat", "已含在 zip 根目录，无需单独拷", "□"],
            ["Node.js 32 位安装包", "仅当中转机确认为 32 位系统时另下", "□"],
            ["手机 + 流量", "联调时临时热点，让中转机访问云平台", "□"],
        ],
        [4, 10.5, 1.2],
    )

    add_heading(doc, "网络说明（必读）", 2)
    add_table(
        doc,
        ["网络类型", "含义", "能否 B1 真北斗", "对策"],
        [
            [
                "完全无网",
                "不能访问 sztu.lingubot.cn",
                "不能验收 B1",
                "联调时段手机热点/临时开外网",
            ],
            [
                "仅内网",
                "北斗↔中转机通，无外网",
                "不能拉真云数据",
                "同上；北斗局域网可先测 register",
            ],
            [
                "内网+外网",
                "能 HTTPS 上云 + 北斗互通",
                "可以",
                "按教程正常联调",
            ],
        ],
        [2.5, 4, 3, 5],
    )
    add_para(
        doc,
        "无网时 CHECK-env 云查询会 FAIL 属正常；START 可能 health=up 但 register/推送会 502。"
        "真 B1 最低要求：联调时段中转机可访问 https://sztu.lingubot.cn 。",
    )

    add_heading(doc, "一、基本信息（到场即填）", 1)
    add_table(
        doc,
        ["项", "填写", "□"],
        [
            ["日期", "", "□"],
            ["我方现场负责人", "", "□"],
            ["中转机 IP", "", "□"],
            ["health 地址", "http://________:8080/health", "□"],
            ["register 地址", "http://________:8080/api/v1/beidou/callback/register", "□"],
            ["北斗回调 URL（register Body 的 url）", "", "□"],
            ["推送 frequency", "4000 ms", "□"],
            ["现场车辆 ID", "LU2606000100", "□"],
            ["北斗联系人 / 电话", "", "□"],
            ["孟泽 / 车端电话", "", "□"],
            ["今日范围", "□ 仅 B1 读推送  □ 含 B2 反控", "□"],
            [
                "中转机网络类型",
                "□ 完全无网  □ 仅内网  □ 内网+外网  □ 联调时临时热点",
                "□",
            ],
        ],
        [4.5, 9.5, 1.5],
    )

    add_heading(doc, "二、到场 10 分钟内必须确认（三方）", 1)
    add_table(
        doc,
        ["#", "确认项", "问谁/怎么测", "结果", "□"],
        [
            ["1", "中转机 IPv4", "现场 IT / ipconfig", "", "□"],
            ["2", "能否访问 sztu.lingubot.cn", "浏览器", "□能  □不能", "□"],
            [
                "2a",
                "若无外网，是否已准备临时上云",
                "手机热点/USB 网卡",
                "□是  □否  □无法解决",
                "□",
            ],
            ["3", "Node.js ≥ 20", "cmd: node -v", "", "□"],
            ["4", "8080 是否可用", "CHECK-env / START", "□空闲  □占用", "□"],
            ["5", "北斗 → 中转机 8080", "北斗访问 health", "□通  □不通", "□"],
            ["6", "中转机 → 北斗回调 URL", "curl 或北斗确认", "□通  □不通", "□"],
            ["7", "北斗回调 URL 已提供", "北斗", "□是  □否", "□"],
            ["8", "孟泽能启 LU2606000100 商用版", "孟泽", "□能  □不能", "□"],
            ["9", "已发 register 地址给北斗", "你", "□是  □否", "□"],
            ["10", "已对齐仅验 B1", "三方", "□是", "□"],
        ],
        [1, 4.5, 4.5, 3.5, 1.2],
    )

    add_heading(doc, "三、程序安装与启动（你方操作）", 1)
    add_table(
        doc,
        ["#", "步骤", "操作", "预期", "□"],
        [
            ["1", "解压部署包", "zip → D:\\beidou-bridge", "有 dist、node_modules", "□"],
            ["2", "放置 .env", "U 盘复制到根目录", "与 dist 同级", "□"],
            ["3", "环境检查", "双击 CHECK-env.bat", "FAIL=0（无网时云项可 FAIL）", "□"],
            ["4", "开防火墙", "管理员 OPEN-firewall-8080.bat", "北斗跨机可访问", "□"],
            ["5", "启动中转", "双击 START-bridge.bat", "窗口保持打开", "□"],
            ["6", "health", "浏览器或 PowerShell", "status = up", "□"],
            ["7", "查云车辆", "双击 QUERY-vehicle.bat", "有 LU2606000100", "□"],
        ],
        [1, 3.5, 4.5, 4.5, 1.2],
    )

    add_heading(doc, "四、联调过程记录（留证）", 1)
    add_table(
        doc,
        ["#", "数据项", "记录", "□"],
        [
            ["A", "CHECK 汇总", "OK=____  FAIL=____", "□"],
            ["B", "health", "status=________", "□"],
            ["C", "云字段 workStatus", "", "□"],
            ["D", "云字段 battery", "", "□"],
            ["E", "云字段 online", "□True  □False", "□"],
            ["F", "云字段 position_xyz", "", "□"],
            ["G", "register 请求 url", "", "□"],
            ["H", "register 响应 code", "", "□"],
            ["I", "register vehicleIds", "须含 LU2606000100", "□"],
            ["J", "推送频率", "约 ____ 秒/次", "□"],
            ["K", "北斗收到 POST", "□是  □否", "□"],
            ["L", "北斗 HTTP 响应 code", "须 1000", "□"],
            ["M", "中转日志 pushed", "□有 LU2606000100", "□"],
        ],
        [1, 4.5, 9, 1.2],
    )

    add_heading(doc, "五、B1 验收结论", 1)
    add_para(doc, "必须全部满足：安装启动 OK、能访问云平台、H code=0、I 含 LU2606000100、K 是、L=1000、M 有。")
    add_para(doc, "不挡 B1：position_xyz 为 0、online 为 False。")
    add_para(doc, "无网阻塞 B1：中转机无法访问 sztu.lingubot.cn 时，不得勾选「B1 总体通过」。")
    add_table(
        doc,
        ["阶段", "日期", "结果", "签字"],
        [
            ["安装与启动", "", "□通过  □未通过", ""],
            ["云平台查询", "", "□通过  □未通过", ""],
            ["北斗 register", "", "□通过  □未通过", ""],
            ["北斗推送 + code:1000", "", "□通过  □未通过", ""],
            ["B1 总体", "", "□通过  □未通过", ""],
        ],
        [4, 3, 4.5, 4],
    )

    add_heading(doc, "六、遗留问题", 1)
    add_table(
        doc,
        ["#", "问题描述", "负责方", "计划处理"],
        [["1", "", "", ""], ["2", "", "", ""], ["3", "", "", ""]],
        [1, 7, 3, 4.5],
    )

    add_heading(doc, "七、给负责人一句话汇报（可复制）", 1)
    add_para(
        doc,
        "中转部署在中转机 ________ :8080，health 正常。云平台可查 LU2606000100（workStatus=____，battery=____）。"
        "北斗 register：□成功 / □失败；vehicleIds 含 LU2606000100。北斗每 4s 收推送并回 code:1000：□是 / □否。"
        "坐标 position_xyz：□非 0 / □仍为 0（不挡 B1）。B2 反控：未测。",
    )

    add_heading(doc, "八、发给北斗的最小信息（到场后补 IP）", 1)
    add_para(doc, "GET  http://<中转机IP>:8080/health", bold=False)
    add_para(doc, "POST http://<中转机IP>:8080/api/v1/beidou/callback/register", bold=False)
    add_para(doc, 'Body: {"url":"<北斗回调URL>","frequency":4000}', bold=False)
    add_para(doc, "收到推送后请 HTTP 响应：{\"code\":1000,\"msg\":\"成功\",\"timestamp\":...}", bold=False)

    add_heading(doc, "九、无网到场操作顺序（仅安装自检）", 1)
    add_table(
        doc,
        ["#", "步骤", "说明", "□"],
        [
            ["1", "U 盘解压 zip", "到 D:\\beidou-bridge", "□"],
            ["2", "复制 .env", "到程序根目录", "□"],
            ["3", "离线安装 Node", "双击 U 盘内 node-v20*.msi", "□"],
            ["4", "CHECK-env.bat", "Node/dist 应 OK；云查询可能 FAIL", "□"],
            ["5", "START-bridge.bat", "health 可能 up；云会报错", "□"],
            ["6", "临时上云", "开热点后再 QUERY / register", "□"],
            ["7", "再验 B1", "有外网后再做第四节记录", "□"],
        ],
        [1, 3.5, 8, 1.2],
    )

    doc.save(OUT)
    print(f"Generated: {OUT}")


if __name__ == "__main__":
    build()
