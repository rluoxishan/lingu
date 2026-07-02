#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate onsite bridge startup & add-vehicle Word guide."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "中转机启动与加车操作手册.docx"
FONT = "微软雅黑"
HEADER_FILL = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 18, True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10, False)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 2")
    available = {s.name for s in doc.styles}
    p = doc.add_paragraph(style=style) if style in available else doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, {1: 14, 2: 12, 3: 11}.get(level, 11), True)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 9.5, False)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], HEADER_FILL)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, 10, True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, 9.5, False)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_title(doc, "中转机 beidou-bridge 启动与加车操作手册")
    add_subtitle(
        doc,
        "MQTT 无外网现场  |  中转机 192.168.199.88  |  现场测试车 LU2605000922  |  版本 2026-07-02",
    )

    add_para(
        doc,
        "适用：中转机完全无外网；车端经 5G/MQTT 上报 1010001；北斗侧 HTTP register + 定时推送。"
        "程序目录示例：D:\\beidou-bridge。",
    )

    add_heading(doc, "第 0 章  现场环境速查（先填/核对）", 1)
    add_table(
        doc,
        ["项", "填写值", "□"],
        [
            ["中转机 IP", "192.168.199.88", "□"],
            ["Mosquitto 端口", "1883（监听 0.0.0.0）", "□"],
            ["bridge HTTP 端口", "8080", "□"],
            ["health", "http://192.168.199.88:8080/health", "□"],
            ["register", "http://192.168.199.88:8080/api/v1/beidou/callback/register", "□"],
            ["北斗回调 URL", "http://192.168.199.89:7055/patroL_vehicle/V1/vehicle/data", "□"],
            ["推送 frequency", "4000 ms", "□"],
            ["现场测试车 SN", "LU2605000922", "□"],
        ],
        [4.5, 9.5, 1.5],
    )

    add_heading(doc, "第 1 章  架构与角色（牢记）", 1)
    add_code(
        doc,
        "5G电脑/真车 ──MQTT:1883──► 中转机(Mosquitto + beidou-bridge:8080) ──HTTP──► 北斗\n"
        "                              ▲\n"
        "                    车发 dev/pub/{SN}  1010001\n"
        "                    车收 dev/sub/{SN}  2010001（反控 B2）",
    )
    add_table(
        doc,
        ["组件", "是否开机自启", "说明"],
        [
            ["Mosquitto", "建议设为自动", "services.msc → Mosquitto Broker"],
            ["beidou-bridge", "默认不自启", "手动双击 START-bridge.bat；窗口勿关"],
            ["北斗 register", "配置持久化", "存 data-site\\beidou-callback.json，重启 bridge 可恢复"],
        ],
        [3.5, 3.5, 8.5],
    )

    add_heading(doc, "第 2 章  首次安装（中转机只做一次）", 1)

    add_heading(doc, "2.1  解压部署包", 2)
    add_para(doc, "1. 将 beidou-bridge-site-20260702.zip 解压到 D:\\beidou-bridge")
    add_para(doc, "2. 确认存在：START-bridge.bat、dist\\main.js、config\\site\\、node_modules\\")

    add_heading(doc, "2.2  安装 Node.js", 2)
    add_para(doc, "1. 双击 node-v24.18.0-x64.msi（或 Node 20 LTS），默认安装")
    add_para(doc, "2. 关闭并重新打开 cmd，执行 node -v，应显示 v20 及以上")

    add_heading(doc, "2.3  安装并配置 Mosquitto", 2)
    add_para(doc, "1. 安装 Mosquitto（U 盘安装包）")
    add_para(doc, "2. 编辑 C:\\Program Files\\mosquitto\\mosquitto.conf，末尾追加：")
    add_code(doc, "listener 1883 0.0.0.0\nallow_anonymous true")
    add_para(doc, "3. Win+R → services.msc → 启动 Mosquitto Broker，启动类型设为「自动」")
    add_para(doc, "4. 验证：netstat -an | findstr 1883  应看到 0.0.0.0:1883")

    add_heading(doc, "2.4  开放防火墙（管理员 CMD，做一次）", 2)
    add_code(
        doc,
        "netsh advfirewall firewall add rule name=\"MQTT-1883\" dir=in action=allow protocol=TCP localport=1883\n"
        "netsh advfirewall firewall add rule name=\"Bridge-8080\" dir=in action=allow protocol=TCP localport=8080",
    )

    add_heading(doc, "2.5  配置现场车辆（见第 3 章）", 2)
    add_para(doc, "编辑 config\\site\\vehicles.yaml，写入现场测试车 LU2605000922。")

    add_heading(doc, "第 3 章  加车辆 / 改车辆 ID（重要）", 1)
    add_para(doc, "不用改程序代码、不用重新编译。只改配置文件，然后重启 bridge 并重新 register。", True)

    add_heading(doc, "3.1  配置文件位置", 2)
    add_code(doc, "D:\\beidou-bridge\\config\\site\\vehicles.yaml")

    add_heading(doc, "3.2  只有一台现场测试车（当前）", 2)
    add_code(
        doc,
        "vehicles:\n"
        "  - vehicleId: \"LU2605000922\"\n"
        "    clientId: \"LU2605000922\"\n"
        "    floor: 1\n"
        "    enabled: true",
    )
    add_table(
        doc,
        ["字段", "含义", "谁提供"],
        [
            ["vehicleId", "推送给北斗的 data.vehicleId", "车端 SN，与 clientId 相同"],
            ["clientId", "MQTT Topic 里的 ID", "必须与车端 SN 完全一致"],
            ["floor", "推北斗时的楼层", "现场约定，一般填 1"],
            ["enabled", "是否启用", "true=订阅并推送；false=忽略该车"],
        ],
        [3, 5.5, 6],
    )
    add_para(doc, "改完必须：① 保存文件  ② 重启 START-bridge.bat  ③ 确认日志 subscribed dev/pub/LU2605000922  ④ 重新 register 北斗", True)

    add_heading(doc, "3.3  多台车同时推北斗（扩展）", 2)
    add_code(
        doc,
        "vehicles:\n"
        "  - vehicleId: \"LU2605000922\"\n"
        "    clientId: \"LU2605000922\"\n"
        "    floor: 1\n"
        "    enabled: true\n"
        "  - vehicleId: \"LU2606000100\"\n"
        "    clientId: \"LU2606000100\"\n"
        "    floor: 1\n"
        "    enabled: true",
    )
    add_para(doc, "每增加或修改一辆车 → 重启 bridge + 重新 register。")

    add_heading(doc, "3.4  发给车端的参数（每台车各一套）", 2)
    add_code(
        doc,
        "【MQTT 连接】\n"
        "Broker IP：192.168.199.88\n"
        "端口：1883\n"
        "完整地址：mqtt://192.168.199.88:1883\n"
        "用户名：（空）\n"
        "密码：（空）\n\n"
        "【设备 ID】\n"
        "SN / clientId：LU2605000922\n\n"
        "【Topic】\n"
        "上报：dev/pub/LU2605000922   （type=1010001，约 1 秒 1 条，QoS 0）\n"
        "订阅：dev/sub/LU2605000922   （收 2010001 下行，QoS 2）\n"
        "回复：dev/reply/LU2605000922 （可选）\n\n"
        "【注意】连的是中转机 Mosquitto，不是云平台公网 EMQX。",
    )

    add_heading(doc, "第 4 章  日常启动步骤（每次联调/开机后）", 1)
    add_para(doc, "按顺序执行，不可跳步：", True)
    add_table(
        doc,
        ["步骤", "操作", "预期/说明"],
        [
            ["①", "确认 Mosquitto 在跑", "netstat 见 0.0.0.0:1883；否则 services.msc 启动"],
            ["②", "双击 START-bridge.bat", "窗口保持打开，勿关"],
            ["③", "curl 检查 health", "status=up, mqttConnected=true"],
            ["④", "确认 MQTT 订阅", "日志：[mqtt] subscribed topic=dev/pub/LU2605000922"],
            ["⑤", "车端/模拟器发 1010001", "真车联调时停掉 5G 模拟脚本"],
            ["⑥", "register 北斗", "首次、改车、改 yaml 后必做"],
            ["⑦", "验收", "日志 pushed vehicle=LU2605000922；北斗回 code:1000"],
        ],
        [1.2, 4.5, 9.8],
    )

    add_heading(doc, "4.1  步骤 ②  启动中转程序", 2)
    add_para(doc, "进入 D:\\beidou-bridge，双击 START-bridge.bat。")
    add_para(doc, "正常日志示例：")
    add_code(
        doc,
        "[OK] MQTT mode, no .env needed\n"
        "[mqtt] connected broker=mqtt://127.0.0.1:1883 ...\n"
        "[mqtt] subscribed topic=dev/pub/LU2605000922\n"
        "beidou-bridge listening on http://0.0.0.0:8080 dataSource=mqtt\n"
        "[scheduler] no registration, push loop idle   ← 未 register 时出现，做完 4.3 即可",
    )

    add_heading(doc, "4.2  步骤 ③  健康检查", 2)
    add_code(doc, "curl http://127.0.0.1:8080/health")
    add_para(doc, "预期：{\"status\":\"up\",\"dataSource\":\"mqtt\",\"mqttConnected\":true}")
    add_para(doc, "北斗侧测试：curl http://192.168.199.88:8080/health")

    add_heading(doc, "4.3  步骤 ⑥  register 北斗", 2)
    add_para(doc, "在中转机 PowerShell 执行（改车或首次必做）：")
    add_code(
        doc,
        '$body = \'{"url":"http://192.168.199.89:7055/patroL_vehicle/V1/vehicle/data","frequency":4000}\'\n'
        "Invoke-RestMethod -Method POST `\n"
        '  -Uri "http://127.0.0.1:8080/api/v1/beidou/callback/register" `\n'
        "  -ContentType \"application/json\" -Body $body",
    )
    add_para(doc, "成功标志：返回 code=0；data.vehicleIds 含 LU2605000922；日志出现 [scheduler] started ...")
    add_para(doc, "给北斗方的 register 地址：")
    add_code(
        doc,
        "POST http://192.168.199.88:8080/api/v1/beidou/callback/register\n\n"
        "Body:\n"
        "{\n"
        '  "url": "http://192.168.199.89:7055/patroL_vehicle/V1/vehicle/data",\n'
        '  "frequency": 4000\n'
        "}",
    )
    add_para(doc, "register 保存在 data-site\\beidou-callback.json。重启 bridge 后一般自动恢复推送，不必重复 register（除非改车或改回调 URL）。")

    add_heading(doc, "4.4  步骤 ⑤  5G 模拟（可选）", 2)
    add_code(
        doc,
        "set VEHICLE_MQTT_BROKER=mqtt://192.168.199.88:1883\n"
        "SIMULATE-vehicle.bat",
    )
    add_para(doc, "模拟器 clientId 必须为 LU2605000922。真车联调时关闭模拟器。")

    add_heading(doc, "第 5 章  停止与重启", 1)
    add_table(
        doc,
        ["操作", "方法"],
        [
            ["停止 bridge", "关闭 START-bridge.bat 窗口或 Ctrl+C"],
            ["停止 Mosquitto", "services.msc → 停止 Mosquitto Broker"],
            ["改 vehicles.yaml 后", "关 bridge → 再开 START-bridge.bat → 重新 register"],
            ["重启电脑后", "Mosquitto 若设自动会自启；bridge 默认需手动启动"],
        ],
        [4, 11.5],
    )

    add_heading(doc, "第 6 章  开机自启（可选，默认未开）", 1)
    add_para(doc, "查是否已注册：Win+R → taskschd.msc → 任务计划程序库 → 找 BeidouBridge")
    add_para(doc, "注册自启（管理员 PowerShell）：")
    add_code(
        doc,
        "cd /d D:\\beidou-bridge\nnpm run autostart:install",
    )
    add_para(
        doc,
        "注意：自启脚本 scripts\\windows\\start-bridge.bat 默认不带 CONFIG_DIR=./config/site。"
        "现场 MQTT 建议继续用手动 START-bridge.bat，或先改自启脚本再加 CONFIG_DIR。",
        True,
    )
    add_para(doc, "取消自启：npm run autostart:remove")

    add_heading(doc, "第 7 章  常见问题", 1)
    add_table(
        doc,
        ["现象", "原因", "处理"],
        [
            ["no registration, push loop idle", "未 register", "执行第 4.3 节 register"],
            ["no status for vehicle=...", "车没发 MQTT 或 ID 不对", "查车端 SN、Topic、5G 网络"],
            ["订阅的是旧 ID", "yaml 未改或未重启", "改 vehicles.yaml → 重启 bridge"],
            ["5G 连不上 1883", "Mosquitto 只监听 127.0.0.1", "conf 加 listener 1883 0.0.0.0"],
            ["北斗 register 失败", "8080 防火墙/bridge 未起", "开防火墙、确认 health"],
            ["push failed", "中转机访问不了北斗", "查 192.168.199.89:7055 网络"],
            ["模拟能通、真车不通", "车端 SN 与 yaml 不一致", "两边统一为 LU2605000922"],
        ],
        [4.5, 4, 7],
    )

    add_heading(doc, "第 8 章  现场速查卡片（可剪贴）", 1)
    add_code(
        doc,
        "【中转机】192.168.199.88\n"
        "【Mosquitto】1883（0.0.0.0，匿名）\n"
        "【bridge 启动】双击 START-bridge.bat（窗口勿关）\n"
        "【health】http://192.168.199.88:8080/health\n"
        "【register】POST http://192.168.199.88:8080/api/v1/beidou/callback/register\n"
        "【北斗回调 url】http://192.168.199.89:7055/patroL_vehicle/V1/vehicle/data\n"
        "【frequency】4000\n"
        "【现场车 SN】LU2605000922\n"
        "【加车配置】config\\site\\vehicles.yaml → 改完重启 + register\n"
        "【车端 Topic】dev/pub/LU2605000922 发 1010001",
    )

    add_heading(doc, "第 9 章  模拟通过后 — 接下来还要做什么", 1)
    add_para(
        doc,
        "若 5G 模拟 + 北斗 register 已验收通过，进入真车联调前请按下列顺序完成。"
        "模拟能通只说明「中转 ↔ MQTT ↔ 北斗 HTTP」链路 OK，真车还需统一 SN 并停掉模拟器。",
        True,
    )

    add_heading(doc, "9.1  必做（真车联调前）", 2)
    add_table(
        doc,
        ["#", "事项", "操作要点", "□"],
        [
            ["1", "中转机改车号", "vehicles.yaml 改为 LU2605000922 → 重启 START-bridge.bat", "□"],
            ["2", "重新 register", "PowerShell POST register；vehicleIds 含 LU2605000922", "□"],
            ["3", "通知北斗", "推送 data.vehicleId 已变为 LU2605000922", "□"],
            ["4", "参数发车端", "IP:1883、SN、Topic（见第 3.4 节）", "□"],
            ["5", "车端确认 SN", "孟泽确认车上 SN = LU2605000922", "□"],
            ["6", "停模拟、上真车", "关 SIMULATE-vehicle.bat；真车发 1010001", "□"],
            ["7", "真车验收", "日志 pushed vehicle=LU2605000922；北斗 code:1000", "□"],
        ],
        [0.8, 2.8, 8.9, 0.8],
    )

    add_heading(doc, "9.2  真车联调当日顺序", 2)
    add_table(
        doc,
        ["顺序", "机器", "动作", "预期"],
        [
            ["①", "中转机", "Mosquitto 运行 + START-bridge.bat 保持", "health 正常；已订阅 dev/pub/LU2605000922"],
            ["②", "5G 电脑", "关闭 SIMULATE-vehicle.bat", "避免同 SN 重复上报"],
            ["③", "真车", "MQTT 连 192.168.199.88:1883，发 1010001", "15 秒内有数据"],
            ["④", "中转机", "看 bridge 日志", "pushed vehicle=LU2605000922"],
            ["⑤", "北斗", "看回调服务", "约 4 秒一条 POST，响应 code:1000"],
        ],
        [1, 2.2, 5.5, 6.8],
    )

    add_heading(doc, "9.3  建议做（稳定运行）", 2)
    add_table(
        doc,
        ["#", "事项", "说明", "□"],
        [
            ["1", "Mosquitto 设自动", "services.msc → Mosquitto Broker → 自动"],
            ["2", "bridge 手动启动", "联调稳定前用手动 START-bridge.bat；稳定后再考虑自启"],
            ["3", "同步 U 盘/现场包", "vehicles.yaml、本手册、START-bridge.bat 拷到中转机"],
            ["4", "验收签字", "填写本章末尾验收表，留档"],
        ],
        [0.8, 3.5, 8.2, 0.8],
    )

    add_heading(doc, "9.4  可选（B1 通过后再做）", 2)
    add_table(
        doc,
        ["项", "说明", "□"],
        [
            ["B2 反控", "北斗 POST navigation → 车收 dev/sub 上 2010001", "□"],
            ["坐标优化", "真车定位正常后 x/y 非 0（不挡 B1）", "□"],
            ["多台车", "vehicles.yaml 加条目 → 重启 → register", "□"],
            ["bridge 开机自启", "npm run autostart:install（需先改自启脚本 CONFIG_DIR）", "□"],
        ],
        [2.5, 10, 1],
    )

    add_heading(doc, "9.5  联调前自检三问", 2)
    add_para(doc, "1. 现场 vehicles.yaml 是否已是 LU2605000922？  □是  □否")
    add_para(doc, "2. 北斗 register 响应 vehicleIds 是否含 LU2605000922？  □是  □否")
    add_para(doc, "3. 车端是否已收到 MQTT 参数并准备联调？  □是  □否")

    add_heading(doc, "第 10 章  明日工作计划（到场填写）", 1)
    add_para(doc, "日期：________    负责人：________    现场联系人（孟泽/北斗）：________")
    add_table(
        doc,
        ["时段", "计划事项", "负责人", "结果", "□"],
        [
            ["上午", "中转机：确认 Mosquitto + bridge 启动；核对 vehicles.yaml", "", "□完成", "□"],
            ["上午", "重新 register；通知北斗 vehicleId=LU2605000922", "", "□完成", "□"],
            ["上午", "微信/当面把车端 MQTT 参数发给孟泽", "", "□完成", "□"],
            ["下午", "停 5G 模拟；真车连 MQTT 发 1010001", "", "□完成", "□"],
            ["下午", "验收：bridge pushed + 北斗 code:1000", "", "□完成", "□"],
            ["下午", "填写验收签字；遗留问题记录", "", "□完成", "□"],
        ],
        [1.5, 5.5, 2, 2.5, 0.8],
    )
    add_para(doc, "遗留问题：")
    add_para(doc, "1. ________________________________________________________________")
    add_para(doc, "2. ________________________________________________________________")

    add_heading(doc, "第 11 章  今日工作记录（摘要模板）", 1)
    add_para(doc, "日期：2026-07-02（示例，到场改写）")
    add_table(
        doc,
        ["类别", "已完成事项"],
        [
            ["架构", "确认无外网现场走 MQTT 模式：车端 5G → 中转机 Mosquitto → bridge → 北斗 HTTP"],
            ["中转机", "Node/Mosquitto/bridge 部署；CHECK-env 通过；START-bridge 运行"],
            ["网络", "Mosquitto 0.0.0.0:1883；防火墙 1883/8080；health mqttConnected:true"],
            ["5G 模拟", "SIMULATE-vehicle 连 192.168.199.88:1883；MQTT 链路打通"],
            ["北斗", "register 成功；北斗侧能收到定时推送（模拟数据）"],
            ["配置", "现场测试车号更正为 LU2605000922（非 LU2606000100）"],
            ["文档", "生成《中转机启动与加车操作手册》Word；明确启动/加车/register 流程"],
            ["待办", "真车联调；中转机 yaml 改车号并 re-register；参数同步孟泽"],
        ],
        [2.5, 12],
    )

    add_heading(doc, "验收签字", 1)
    add_table(
        doc,
        ["阶段", "日期", "结果", "签字"],
        [
            ["Mosquitto + 防火墙", "", "□通过 □未通过", ""],
            ["bridge 启动 + health", "", "□通过 □未通过", ""],
            ["车辆 yaml + 车端 MQTT", "", "□通过 □未通过", ""],
            ["北斗 register", "", "□通过 □未通过", ""],
            ["5G 模拟验收", "", "□通过 □未通过", ""],
            ["真车 MQTT 上报", "", "□通过 □未通过", ""],
            ["真车推送北斗", "", "□通过 □未通过", ""],
        ],
        [4.5, 2.5, 4, 4.5],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
