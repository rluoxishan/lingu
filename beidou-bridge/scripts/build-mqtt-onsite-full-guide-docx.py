#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate full MQTT onsite deployment & troubleshooting Word guide."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "MQTT现场部署与故障排查手册.docx"
FONT = "微软雅黑"
HEADER_FILL = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, size=10.5, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 18, True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10, False)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    style = style_map.get(level, "Heading 2")
    available = {s.name for s in doc.styles}
    p = doc.add_paragraph(style=style) if style in available else doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, {1: 14, 2: 12, 3: 11}.get(level, 11), True)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 9.5, False)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], HEADER_FILL)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, 10, True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, 9.5, False)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_title(doc, "beidou-bridge MQTT 现场部署与故障排查手册")
    add_subtitle(
        doc,
        "架构：5G电脑/模拟车 ──MQTT──► 中转机 ──HTTP──► 北斗  |  现场车：LU2606000100  |  版本：2026-07-02",
    )

    add_para(
        doc,
        "用途：进机房部署中转机（完全无外网）、5G 模拟电脑联调、接真北斗验收。"
        "中转机不访问云平台，不需要 .env 云账号。",
    )

    add_heading(doc, "一、架构与验收目标", 1)
    add_para(doc, "数据链路：", bold=True)
    add_code(
        doc,
        "5G电脑/真车  --MQTT 1010001 (1883, dev/pub/LU2606000100)-->  中转机(Mosquitto+bridge:8080)  --HTTP POST-->  北斗回调",
    )
    add_table(
        doc,
        ["段", "协议", "谁部署"],
        [
            ["车 → 中转", "MQTT 1010001", "5G 电脑 SIMULATE-vehicle.bat 或真车"],
            ["中转 → 北斗", "HTTP register + 定时 POST", "中转机 START-bridge.bat"],
            ["中转 ↔ 云", "无", "不需要"],
        ],
        [3, 5, 7.5],
    )
    add_para(doc, "B1 验收标准：北斗 register 成功；周期性收到 POST；HTTP 响应 code=1000；含 vehicleId=LU2606000100。", bold=True)

    add_heading(doc, "二、出发前 U 盘必备（有网电脑准备）", 1)
    add_table(
        doc,
        ["#", "文件", "说明", "□已拷"],
        [
            ["1", "beidou-bridge-site-20260702.zip", "MQTT 版部署包（勿用 20260701 旧包）", "□"],
            ["2", "node-v20 或 v24 *.msi", "中转机 + 5G 电脑装 Node", "□"],
            ["3", "mosquitto-*-install-windows-x64.exe", "仅中转机安装 MQTT Broker", "□"],
            ["4", "本文档 + 给北斗方-接口一页纸.txt", "打印或 U 盘携带", "□"],
        ],
        [0.8, 5, 7.7, 1.2],
    )
    add_para(doc, "不需要：.env、云账号、在中转机 npm install（zip 已含 node_modules）。", bold=True)

    add_heading(doc, "三、环境信息表（到场填写）", 1)
    add_table(
        doc,
        ["项", "填写", "□"],
        [
            ["日期", "", "□"],
            ["中转机 IP", "", "□"],
            ["5G 电脑 IP", "", "□"],
            ["5G 能否 ping 通中转机", "□能  □不能", "□"],
            ["health", "http://____:8080/health", "□"],
            ["register", "http://____:8080/api/v1/beidou/callback/register", "□"],
            ["北斗回调 URL（register Body 的 url）", "", "□"],
            ["推送 frequency", "4000 ms", "□"],
            ["北斗联系人", "", "□"],
        ],
        [4.5, 9.5, 1.5],
    )

    add_heading(doc, "四、阶段 A：5G 模拟电脑部署", 1)
    add_heading(doc, "4.1 拷贝程序", 2)
    add_table(
        doc,
        ["步骤", "操作", "预期", "□"],
        [
            ["A1", "U 盘文件夹 beidou-bridge-site-20260702 整夹复制到 D:\\beidou-bridge", "含 node_modules、dist", "□"],
            ["A2", "双击 node-v24 或 node-v20 *.msi 安装 Node", "默认下一步", "□"],
            ["A3", "关闭并重开 cmd，执行 node -v", "显示 v20+ 或 v24", "□"],
            ["A4", "cd D:\\beidou-bridge 后双击 CHECK-5g-sim.bat", "全部 [OK]", "□"],
        ],
        [1.2, 8.5, 4.8, 1.2],
    )
    add_para(doc, "注意：中转机 Mosquitto 未就绪前，不要运行 SIMULATE-vehicle.bat。", bold=True)

    add_heading(doc, "五、阶段 B：中转机部署（机房 · 完全无网）", 1)
    add_heading(doc, "5.1 安装 Node.js", 2)
    add_table(
        doc,
        ["步骤", "操作", "预期", "□"],
        [
            ["B1", "U 盘安装 node *.msi", "安装完成", "□"],
            ["B2", "cmd: node -v", "v20+ 或 v24", "□"],
        ],
        [1.2, 8.5, 4.8, 1.2],
    )

    add_heading(doc, "5.2 安装 Mosquitto（必装）", 2)
    add_table(
        doc,
        ["步骤", "操作", "预期", "□"],
        [
            ["B3", "安装 mosquitto-*-windows-x64.exe", "默认路径", "□"],
            ["B4", "编辑 C:\\Program Files\\mosquitto\\mosquitto.conf 末尾追加两行", "见下方配置", "□"],
            ["B5", "services.msc 启动 Mosquitto Broker，设为自动", "服务运行中", "□"],
            ["B6", "cmd: netstat -an | findstr :1883", "LISTENING", "□"],
        ],
        [1.2, 8.5, 4.8, 1.2],
    )
    add_para(doc, "mosquitto.conf 末尾追加：", bold=True)
    add_code(doc, "listener 1883 0.0.0.0")
    add_code(doc, "allow_anonymous true")

    add_heading(doc, "5.3 部署中转程序", 2)
    add_table(
        doc,
        ["步骤", "操作", "预期", "□"],
        [
            ["B7", "复制 beidou-bridge-site-20260702 到 D:\\beidou-bridge", "有 dist\\main.js", "□"],
            ["B8", "管理员运行 OPEN-firewall-1883.bat", "防火墙规则已加", "□"],
            ["B9", "管理员运行 OPEN-firewall-8080.bat", "北斗可访问 8080", "□"],
            ["B10", "cmd: ipconfig 记录 IPv4", "填入第三节环境表", "□"],
            ["B11", "双击 CHECK-env.bat", "summary FAIL=0", "□"],
        ],
        [1.2, 8.5, 4.8, 1.2],
    )

    add_heading(doc, "六、阶段 C：双机 Mock 联调（接北斗前必做）", 1)
    add_table(
        doc,
        ["步骤", "在哪", "操作", "成功标志", "□"],
        [
            ["C1", "5G 电脑", "ping <中转机IP>", "有回复", "□"],
            ["C2", "中转机", "双击 RUN-mqtt-lab-relay.bat", "开 Mock+bridge 两窗口", "□"],
            ["C3", "中转机", "curl http://127.0.0.1:8080/health", "mqttConnected:true", "□"],
            ["C4", "5G 电脑", "set VEHICLE_MQTT_BROKER=mqtt://<中转机IP>:1883", "环境变量设好", "□"],
            ["C5", "5G 电脑", "SIMULATE-vehicle.bat", "[sim] MQTT 已连接", "□"],
            ["C6", "中转机 bridge", "看日志", "pushed vehicle=LU2606000100", "□"],
            ["C7", "Mock 北斗窗口", "看输出", "PUSH #1 #2…", "□"],
        ],
        [1, 2, 5.5, 5.8, 1.2],
    )
    add_para(doc, "Mock 通过 = 5G → 中转机 → HTTP 推送 软件链路 OK。", bold=True)

    add_heading(doc, "七、阶段 D：接真北斗", 1)
    add_table(
        doc,
        ["步骤", "操作", "说明", "□"],
        [
            ["D1", "关闭 Mock 北斗窗口", "保留 Mosquitto 服务", "□"],
            ["D2", "双击 START-bridge.bat", "窗口勿关", "□"],
            ["D3", "北斗 POST register", "Body: url+frequency:4000", "□"],
            ["D4", "5G 电脑继续 SIMULATE-vehicle.bat", "或真车发 1010001", "□"],
            ["D5", "北斗确认 POST + code:1000", "B1 验收", "□"],
        ],
        [1.2, 5, 8.3, 1.2],
    )
    add_para(doc, "发给北斗的 register 示例：", bold=True)
    add_code(doc, "POST http://<中转机IP>:8080/api/v1/beidou/callback/register")
    add_code(doc, 'Body: {"url":"<北斗回调URL>","frequency":4000}')

    add_heading(doc, "八、日常启动顺序", 1)
    add_table(
        doc,
        ["顺序", "动作", "□"],
        [
            ["1", "Mosquitto 服务（1883）", "□"],
            ["2", "START-bridge.bat（8080）", "□"],
            ["3", "5G 电脑 SIMULATE 或真车 MQTT", "□"],
            ["4", "北斗 register（改 URL 时重做）", "□"],
        ],
        [1.5, 12, 1.2],
    )

    add_heading(doc, "九、故障排查（现象 → 原因 → 解决步骤）", 1)

    add_heading(doc, "9.1 网络与 MQTT", 2)
    add_table(
        doc,
        ["现象", "可能原因", "解决步骤"],
        [
            [
                "5G 电脑 SIMULATE 报连接失败",
                "网络不通 / 防火墙 / Mosquitto 未起",
                "① ping 中转机 IP ② 中转机 netstat 查 1883 ③ 管理员 OPEN-firewall-1883.bat ④ 确认 Mosquitto 服务已启动",
            ],
            [
                "health 里 mqttConnected:false",
                "Mosquitto 未运行或 mqtt.yaml 错误",
                "① 启动 Mosquitto ② 确认 config\\site\\mqtt.yaml 为 brokerUrl: mqtt://127.0.0.1:1883 ③ 重启 START-bridge.bat",
            ],
            [
                "ping 不通中转机",
                "5G 模块/网线/网段不对",
                "① 确认 5G 与中转机同网段 ② 联系现场网络/集成商 ③ 网络通后再测 SIMULATE",
            ],
        ],
        [3.5, 3.5, 8.5],
    )

    add_heading(doc, "9.2 中转程序", 2)
    add_table(
        doc,
        ["现象", "可能原因", "解决步骤"],
        [
            [
                "CHECK-env FAIL=dist 缺失",
                "zip 不完整或拷错目录",
                "重新从 U 盘完整复制 beidou-bridge-site-20260702 文件夹",
            ],
            [
                "CHECK-env FAIL=1883",
                "Mosquitto 未监听",
                "按 5.2 节配置并启动服务，netstat 确认 LISTENING",
            ],
            [
                "START-bridge 一闪退出",
                "无 Node 或无 dist",
                "node -v 检查；确认 D:\\beidou-bridge\\dist\\main.js 存在",
            ],
            [
                "register 成功但无 PUSH",
                "无 1010001 上报",
                "5G 电脑运行 SIMULATE-vehicle.bat；或真车 MQTT 指向中转机 IP:1883",
            ],
            [
                "日志 no status for vehicle",
                "15 秒内未收到 MQTT",
                "检查 Topic 是否为 dev/pub/LU2606000100；clientId 是否为 LU2606000100",
            ],
            [
                "PUSH 里 x/y=0",
                "坐标字段未上报",
                "模拟器/真车应发 position_xyz；config 中 positionMode: map_xy",
            ],
        ],
        [3.5, 3.5, 8.5],
    )

    add_heading(doc, "9.3 北斗 HTTP", 2)
    add_table(
        doc,
        ["现象", "可能原因", "解决步骤"],
        [
            [
                "北斗 register 连不上",
                "8080 防火墙或 IP 错",
                "① OPEN-firewall-8080.bat ② 北斗侧 ping/curl health ③ 确认 register URL 中 IP 正确",
            ],
            [
                "北斗收不到 POST",
                "register 的 url 错或 bridge 未跑",
                "① START-bridge 窗口保持打开 ② 核对 register Body 中 url ③ 看 bridge 是否有 pushed 日志",
            ],
            [
                "北斗返回非 1000",
                "北斗侧校验失败",
                "把 POST body 给北斗方核对字段；对照 1010001 映射文档",
            ],
        ],
        [3.5, 3.5, 8.5],
    )

    add_heading(doc, "9.4 常见误区", 2)
    add_table(
        doc,
        ["误区", "正确做法"],
        [
            ["用 20260701 旧 zip", "必须用 20260702（含 MQTT 代码）"],
            ["中转机需要 .env 云账号", "MQTT 模式不需要"],
            ["在中转机 npm install", "zip 已含 node_modules，无网不要 npm"],
            ["未 Mock 直接接北斗", "先 RUN-mqtt-lab-relay + SIMULATE 验证双机"],
            ["QUERY-vehicle.bat 查云", "无外网无用，看 health 和 bridge 日志"],
        ],
        [5, 10.5],
    )

    add_heading(doc, "十、开发机本机预演（可选）", 1)
    add_para(doc, "在有网开发机验证程序：cd 到 beidou-bridge 后执行 npm run test:mqtt-lab，应显示 PASS。")
    add_para(doc, "本机预演不能代替 5G↔中转机 网络测试，但可确认程序无问题。")

    add_heading(doc, "十一、验收签字", 1)
    add_table(
        doc,
        ["阶段", "日期", "结果", "签字"],
        [
            ["5G 电脑 CHECK-5g-sim", "", "□通过  □未通过", ""],
            ["中转机 CHECK-env", "", "□通过  □未通过", ""],
            ["双机 Mock PUSH", "", "□通过  □未通过", ""],
            ["真北斗 register+推送", "", "□通过  □未通过", ""],
            ["B1 总体", "", "□通过  □未通过", ""],
        ],
        [4, 3, 4.5, 4],
    )

    doc.save(OUT)
    print(f"Generated: {OUT}")


if __name__ == "__main__":
    build()
